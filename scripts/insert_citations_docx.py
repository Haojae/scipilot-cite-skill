"""
SciPilot-cite :: insert_citations_docx.py
Word 文档引用插入引擎
保持原文格式不变，仅添加引用标记和 References 章节
"""
from __future__ import annotations

import argparse
import json
import os
import re
import sys
from typing import Any

from format_citation import format_citation
from utils import assign_citation_numbers, extract_keywords, logger

try:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError as e:
    print(
        "[SciPilot-cite] python-docx is required: pip install python-docx",
        file=sys.stderr,
    )
    raise


SENTENCE_END_RE = re.compile(r"(?<=[.!?。！？])\s+(?=\S)")
HEADING_STYLE_RE = re.compile(r"Heading\s*(\d+)", re.I)


def parse_docx(docx_path: str) -> dict[str, Any]:
    """读取 .docx 并提取段落、章节结构、现有引用。"""
    doc = Document(docx_path)
    paragraphs = []
    sections: list[dict] = []
    current_section = None

    for i, para in enumerate(doc.paragraphs):
        text = para.text
        style_name = para.style.name if para.style else ""
        heading_match = HEADING_STYLE_RE.match(style_name)
        is_heading = bool(heading_match) or style_name.lower().startswith("title")
        para_record = {
            "index": i,
            "text": text,
            "style": style_name,
            "is_heading": is_heading,
            "heading_level": int(heading_match.group(1)) if heading_match else 0,
        }
        paragraphs.append(para_record)

        if is_heading and text.strip():
            if current_section:
                current_section["end_index"] = i - 1
                sections.append(current_section)
            current_section = {"name": text.strip(), "start_index": i + 1, "end_index": len(doc.paragraphs) - 1}

    if current_section:
        sections.append(current_section)

    existing_refs = []
    for sec in sections:
        if any(k in sec["name"].lower() for k in ("reference", "bibliography", "参考文献")):
            for p in paragraphs[sec["start_index"] : sec["end_index"] + 1]:
                if p["text"].strip():
                    existing_refs.append(p["text"])

    return {
        "doc": doc,
        "paragraphs": paragraphs,
        "sections": sections,
        "existing_references": existing_refs,
        "path": docx_path,
    }


def _section_priority(name: str) -> int:
    n = name.lower()
    if "related" in n or "background" in n or "literature" in n:
        return 1
    if "introduction" in n or "intro" in n or "引言" in n:
        return 2
    if "method" in n or "approach" in n or "model" in n or "方法" in n:
        return 3
    if "experiment" in n or "result" in n or "evaluation" in n or "实验" in n:
        return 4
    if "discussion" in n or "conclu" in n or "讨论" in n or "结论" in n:
        return 5
    return 9


def find_insertion_points(parsed: dict, papers: list[dict]) -> list[dict]:
    """为每篇文献选择最相关的段落作为插入点。"""
    sections = parsed["sections"]
    paragraphs = parsed["paragraphs"]
    if not sections:
        eligible_idx = [p["index"] for p in paragraphs if p["text"].strip() and not p["is_heading"]]
        if not eligible_idx:
            return []
        sections = [{"name": "body", "start_index": eligible_idx[0], "end_index": eligible_idx[-1]}]

    sections = [s for s in sections if not any(k in s["name"].lower() for k in ("reference", "bibliography", "参考文献"))]

    used_para_idx: set[int] = set()
    plan: list[dict] = []

    for idx, paper in enumerate(papers):
        kws = set(
            extract_keywords(
                (paper.get("title", "") + " " + paper.get("abstract", "")).lower(), top_n=8
            )
        )
        best = None
        best_score = -1.0
        for sec in sections:
            base = 100 - _section_priority(sec["name"]) * 5
            for p in paragraphs[sec["start_index"] : sec["end_index"] + 1]:
                if p["is_heading"] or not p["text"].strip():
                    continue
                if p["index"] in used_para_idx:
                    continue
                lower = p["text"].lower()
                overlap = sum(1 for k in kws if k in lower)
                length_bonus = min(len(p["text"]) / 200, 2.0)
                score = base + overlap * 10 + length_bonus
                if score > best_score:
                    best_score = score
                    best = (sec["name"], p["index"], p["text"])
        if not best:
            continue
        sec_name, para_idx, _text = best
        used_para_idx.add(para_idx)
        plan.append(
            {
                "paper_idx": idx,
                "section_name": sec_name,
                "position": para_idx,
                "paragraph_index": para_idx,
            }
        )

    return assign_citation_numbers(plan)


def insert_citations_to_docx(
    docx_path: str,
    papers: list[dict],
    style: str,
    insertion_plan: list[dict],
    output_path: str,
) -> dict:
    """
    在 Word 中插入 [N] 标记并在文末添加 References 章节。
    返回 {ok, output, insertions, papers}。
    """
    parsed = parse_docx(docx_path)
    doc = parsed["doc"]

    plan_by_para: dict[int, list[dict]] = {}
    for item in insertion_plan:
        plan_by_para.setdefault(item["paragraph_index"], []).append(item)

    for para_idx, items in plan_by_para.items():
        para = doc.paragraphs[para_idx]
        items.sort(key=lambda x: x["citation_number"])
        nums = [it["citation_number"] for it in items]
        if len(nums) > 1:
            cite_marker = " [" + ",".join(str(n) for n in nums) + "]"
        else:
            cite_marker = f" [{nums[0]}]"

        runs = para.runs
        last_run = runs[-1] if runs else None
        if last_run:
            txt = last_run.text or ""
            stripped = txt.rstrip()
            trailing = txt[len(stripped):]
            if stripped.endswith((".", "!", "?", "。", "！", "？")):
                last_run.text = stripped[:-1] + cite_marker + stripped[-1] + trailing
            else:
                last_run.text = stripped + cite_marker + trailing
        else:
            para.add_run(cite_marker)

    papers_by_number: dict[int, dict] = {}
    for item in insertion_plan:
        papers_by_number[item["citation_number"]] = papers[item["paper_idx"]]

    has_ref_section = any("reference" in s["name"].lower() or "参考文献" in s["name"] for s in parsed["sections"])

    if not has_ref_section:
        doc.add_paragraph()
        heading = doc.add_paragraph("References")
        try:
            heading.style = doc.styles["Heading 1"]
        except KeyError:
            heading.runs[0].bold = True

    for n in sorted(papers_by_number.keys()):
        paper = papers_by_number[n]
        entry = format_citation(paper, style, number=n)
        doc.add_paragraph(entry)

    doc.save(output_path)
    logger.info(f"Wrote {output_path} with {len(insertion_plan)} citations across {len(papers_by_number)} papers")
    return {
        "ok": True,
        "output": output_path,
        "insertions": len(insertion_plan),
        "papers": len(papers_by_number),
    }


def run(docx_path: str, papers_json_path: str, style: str = "ieee", output_path: str | None = None) -> dict:
    with open(papers_json_path, encoding="utf-8") as f:
        papers = json.load(f)
    parsed = parse_docx(docx_path)
    plan = find_insertion_points(parsed, papers)
    if not plan:
        logger.error("No insertion points found")
        return {"ok": False, "reason": "no insertion points"}
    out_path = output_path or docx_path.replace(".docx", "_scipilot.docx")
    return insert_citations_to_docx(docx_path, papers, style, plan, out_path)


def _cli() -> int:
    p = argparse.ArgumentParser(description="SciPilot-cite Word citation inserter")
    p.add_argument("docx_file")
    p.add_argument("papers_json")
    p.add_argument("--style", default="ieee")
    p.add_argument("--output")
    args = p.parse_args()
    r = run(args.docx_file, args.papers_json, args.style, args.output)
    print(json.dumps(r, ensure_ascii=False, indent=2))
    return 0 if r.get("ok") else 1


if __name__ == "__main__":
    sys.exit(_cli())
